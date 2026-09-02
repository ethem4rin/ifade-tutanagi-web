# -*- coding: utf-8 -*-
"""İfade Tutanağı .docx üretici.

Sayfa düzeni 'deneme ifade.docx' örneğinden birebir çıkarılmıştır:
  - A4, kenar boşlukları sol 2.0 / sağ 1.5 / üst 1.0 / alt 1.5 cm
  - Antet: amblem (2.22 cm, ortalı) + T.C. / Bakanlık / Başkanlık satırları
  - Üst bilgi satırları sekme ile, iki nokta 3.75 cm'de hizalı
  - Görev paragrafı: iki yana yaslı, ilk satır girintisi 1.25 cm
  - Kimlik tablosu: kenarlıksız, 6 sütun, son satır birleşik
  - Gövde paragrafları: iki yana yaslı, ilk satır 1.0 cm, önce/sonra 6 pt
  - İmza tablosu: sayfa genişliği, 4 sütun, ortalı
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import ifade_metni


def _kaynak(goreli: str) -> str:
    """Kaynak dosya yolu — hem geliştirmede hem PyInstaller exe'sinde çalışır."""
    if getattr(sys, "frozen", False):
        base = sys._MEIPASS
    else:
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, goreli)


_AMBLEM = _kaynak(os.path.join("sablonlar", "amblem.png"))
_YAZI = "Times New Roman"
_BOYUT = 12          # docDefaults: 24 half-point = 12 pt
_KOLON = Cm(3.75)    # üst bilgi satırlarında ":" hizası

# Antet metinleri
_ANTET = ["T.C.", "ÇALIŞMA VE SOSYAL GÜVENLİK BAKANLIĞI", "Rehberlik ve Teftiş Başkanlığı"]

# Kimlik tablosu sütun genişlikleri (twips, örnekten)
_KIMLIK_SUTUN = [1696, 142, 2984, 1694, 142, 3362]
# Antet tablosu sütun genişlikleri (twips, örnekten)
_ANTET_SUTUN = [1602, 7170, 1550]
# İmza tablosu sütun genişlikleri (twips, örnekten)
_IMZA_SUTUN = [2693, 2695, 1702, 2831]


# --------------------------------------------------------------------------
# yardımcılar
# --------------------------------------------------------------------------
def _font(run, boyut=_BOYUT, kalin=False):
    run.font.name = _YAZI
    run.font.size = Pt(boyut)
    run.font.bold = kalin
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), _YAZI)
    return run


def _ekle(p, metin, boyut=_BOYUT, kalin=False):
    return _font(p.add_run(metin), boyut, kalin)


def _p(doc, hiza=None, once=None, sonra=0, ilk_girinti=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if hiza is not None:
        p.alignment = hiza
    pf.space_before = Pt(once) if once is not None else Pt(0)
    pf.space_after = Pt(sonra) if sonra is not None else Pt(0)
    if ilk_girinti is not None:
        pf.first_line_indent = ilk_girinti
    return p


def _tablo_kenarsiz(tablo):
    """Tüm kenarlıkları kaldırır (örnekteki iki tablo da kenarlıksız)."""
    tblPr = tablo._tbl.tblPr
    for eski in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(eski)
    borders = OxmlElement("w:tblBorders")
    for kenar in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + kenar)
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        borders.append(e)
    tblPr.append(borders)


def _hucre(hucre, metin, hiza=WD_ALIGN_PARAGRAPH.JUSTIFY, kalin=False):
    hucre.text = ""
    p = hucre.paragraphs[0]
    p.alignment = hiza
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _ekle(p, metin, _BOYUT, kalin)
    return p


def _sutun_genislik(tablo, twips_list):
    """Sütun genişliklerini sabitler (tblLayout=fixed + hücre genişlikleri)."""
    tblPr = tablo._tbl.tblPr
    # sabit yerleşim: Word sütunları içeriğe göre daraltıp genişletmesin
    for eski in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(eski)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    # toplam genişlik
    for eski in tblPr.findall(qn("w:tblW")):
        tblPr.remove(eski)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), str(sum(twips_list)))
    tw.set(qn("w:type"), "dxa")
    tblPr.append(tw)
    # hücre iç boşluğu sıfır (örnekte böyle; etiketler alt satıra kaymasın)
    for eski_mar in tblPr.findall(qn("w:tblCellMar")):
        tblPr.remove(eski_mar)
    mar = OxmlElement("w:tblCellMar")
    for yan in ("left", "right"):
        e = OxmlElement("w:" + yan)
        e.set(qn("w:w"), "0")
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tblPr.append(mar)
    # tblGrid sütunları
    grid = tablo._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, w in zip(grid.findall(qn("w:gridCol")), twips_list):
            gc.set(qn("w:w"), str(w))
    # hücre genişlikleri (birleşik hücreleri bozmadan)
    for satir in tablo.rows:
        hucreler = satir.cells
        gorulen = []
        for i, hucre in enumerate(hucreler):
            if hucre._tc in gorulen:
                continue
            gorulen.append(hucre._tc)
            span = sum(1 for h in hucreler if h._tc is hucre._tc)
            if i < len(twips_list):
                hucre.width = Twips(sum(twips_list[i:i + span]))


# --------------------------------------------------------------------------
def _sayfa_ayar(doc):
    s = doc.sections[0]
    s.page_width = Twips(11906)     # A4
    s.page_height = Twips(16838)
    s.top_margin = Twips(567)       # 1.0 cm
    s.right_margin = Twips(851)     # 1.5 cm
    s.bottom_margin = Twips(851)    # 1.5 cm
    s.left_margin = Twips(1134)     # 2.0 cm
    s.header_distance = Twips(570)
    s.footer_distance = Twips(709)

    normal = doc.styles["Normal"]
    normal.font.name = _YAZI
    normal.font.size = Pt(_BOYUT)
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), _YAZI)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def _antet(doc):
    """Antet: 1x3 kenarlıksız tablo -> [amblem] [T.C./Bakanlık/Başkanlık] [boş].

    Örnekteki header birebir: tablo 10322 twips, sütunlar 1602/7170/1550,
    ortalanmış, metinler kalın ve ortalı (Başlık 4 stili)."""
    hdr = doc.sections[0].header
    hdr.is_linked_to_previous = False

    t = hdr.add_table(rows=1, cols=3, width=Twips(_ANTET_SUTUN[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    _tablo_kenarsiz(t)

    # sol hücre: amblem
    c0 = t.cell(0, 0)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.right_indent = Twips(176)
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(0)
    if os.path.exists(_AMBLEM):
        p0.add_run().add_picture(_AMBLEM, width=Cm(2.22), height=Cm(2.22))

    # orta hücre: antet metinleri (kalın, ortalı)
    c1 = t.cell(0, 1)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c1.text = ""
    for i, metin in enumerate(_ANTET):
        p = c1.paragraphs[0] if i == 0 else c1.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.left_indent = Twips(342)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _ekle(p, metin, kalin=True)

    _sutun_genislik(t, _ANTET_SUTUN)

    # örnekteki gibi tablodan sonra boş satır
    bp = hdr.add_paragraph()
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after = Pt(0)


def _ust_bilgi(doc, etiket, deger):
    """'Sicil No.' + sekme + ':  değer' — iki nokta 3.75 cm'de hizalı."""
    p = _p(doc, sonra=0)
    p.paragraph_format.tab_stops.add_tab_stop(_KOLON, WD_TAB_ALIGNMENT.LEFT)
    _ekle(p, etiket)
    _ekle(p, "\t:  " + str(deger))
    return p


# --------------------------------------------------------------------------
def uret(isyeri: dict, baslik: dict, isci: dict, sira: int, kayit_yolu: str,
         govde_list=None, kapanis_metni=None) -> str:
    """Bir işçi için ifade tutanağı .docx üretir, dosya yolunu döndürür.

    govde_list / kapanis_metni verilirse (önizlemede elle düzenlenmiş metin)
    ifade_metni yerine onlar kullanılır."""
    doc = Document()
    _sayfa_ayar(doc)
    _antet(doc)

    # --- başlık (ortalı, kalın) ---
    p = _p(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _ekle(p, "İFADE TUTANAĞI – " + str(sira) + "-", kalin=True)
    _p(doc)                                   # boş satır

    # --- üst bilgi satırları (sekmeli) ---
    _ust_bilgi(doc, "Sicil No.", isyeri.get("sicil", ""))
    _ust_bilgi(doc, "İşyerinin Unvanı", isyeri.get("unvan_kisa", ""))
    _ust_bilgi(doc, "İşveren", isyeri.get("isveren", ""))
    _ust_bilgi(doc, "İşyerinin Adresi", isyeri.get("adres", ""))
    _p(doc)                                   # boş satır

    # --- görev yazısı paragrafı ---
    gp = _p(doc, WD_ALIGN_PARAGRAPH.JUSTIFY, ilk_girinti=Twips(708))
    _ekle(gp,
          "Çalışma ve Sosyal Güvenlik Bakanlığı İstanbul Rehberlik Ve Teftiş Grup "
          "Başkanlığı’nın " + str(baslik.get("tarih", "")) + " tarihli ve "
          + str(baslik.get("sayi_son", "")) + " (" + str(isyeri.get("gorev_no", ""))
          + ") sayılı görev yazısı uyarınca yapılan inceleme teftişinde, "
          "aşağıda kimliği tespit edilen kişinin ifadesine başvuruldu:")
    _p(doc)                                   # boş satır

    # --- kimlik tablosu (kenarlıksız) ---
    kim = doc.add_table(rows=4, cols=6)
    kim.alignment = WD_TABLE_ALIGNMENT.LEFT
    kim.autofit = False
    _tablo_kenarsiz(kim)
    satirlar = [
        ("Adı Soyadı", isci.get("ad_soyad", ""), "Doğum Yeri/Yılı", isci.get("dogum_yeri_yili", "")),
        ("Baba/Ana Adı", isci.get("baba_ana", ""), "Görevi", isci.get("gorevi", "")),
        ("T.C. Kimlik No", isci.get("tc", ""), "Telefon No", isci.get("telefon", "")),
    ]
    for i, (e1, d1, e2, d2) in enumerate(satirlar):
        _hucre(kim.cell(i, 0), e1)
        _hucre(kim.cell(i, 1), ":")
        _hucre(kim.cell(i, 2), d1)
        _hucre(kim.cell(i, 3), e2)
        _hucre(kim.cell(i, 4), ":")
        _hucre(kim.cell(i, 5), d2)
    # son satır: İkametgahı — kalan sütunlar birleşik
    _hucre(kim.cell(3, 0), "İkametgahı")
    _hucre(kim.cell(3, 1), ":")
    birlesik = kim.cell(3, 2).merge(kim.cell(3, 5))
    _hucre(birlesik, isci.get("ikametgah", ""))
    _sutun_genislik(kim, _KIMLIK_SUTUN)

    # --- ifade gövdesi (ilk satır 1.0 cm, önce/sonra 6 pt) ---
    paragraflar = govde_list if govde_list is not None else ifade_metni.govde(isci)
    for metin in paragraflar:
        if not str(metin).strip():
            continue
        bp = _p(doc, WD_ALIGN_PARAGRAPH.JUSTIFY, once=6, sonra=6, ilk_girinti=Twips(567))
        _ekle(bp, str(metin))

    # --- kapanış (ilk satır 1.25 cm) ---
    kp = _p(doc, WD_ALIGN_PARAGRAPH.JUSTIFY, ilk_girinti=Twips(708))
    _ekle(kp, kapanis_metni if kapanis_metni is not None else ifade_metni.kapanis(isci))
    _p(doc, WD_ALIGN_PARAGRAPH.JUSTIFY, ilk_girinti=Twips(708))   # örnekteki iki boş satır
    _p(doc, WD_ALIGN_PARAGRAPH.JUSTIFY, ilk_girinti=Twips(708))

    # --- imza tablosu ---
    imz = doc.add_table(rows=2, cols=4)
    imz.autofit = False
    _tablo_kenarsiz(imz)
    imzalar = [
        (baslik.get("kidemli", ""), "İş Müfettişi"),
        (baslik.get("kidemsiz", ""), "İş Müfettişi"),
        ("", ""),
        (isci.get("ad_soyad", ""), "İfade Sahibi"),
    ]
    for j, (ad, rol) in enumerate(imzalar):
        _hucre(imz.cell(0, j), ad, hiza=WD_ALIGN_PARAGRAPH.CENTER)
        _hucre(imz.cell(1, j), rol, hiza=WD_ALIGN_PARAGRAPH.CENTER)
    _sutun_genislik(imz, _IMZA_SUTUN)

    os.makedirs(os.path.dirname(os.path.abspath(kayit_yolu)), exist_ok=True)
    doc.save(kayit_yolu)
    return kayit_yolu
